
import PyPDF2
pdfFile = open('PREPRI STATEMENT 20210312.pdf', 'rb') # read binary
reader = PyPDF2.PdfReader(pdfFile)

print(len(reader.pages))

for page in range(len(reader.pages)):
    print(reader.pages[page].extract_text())

writer = PyPDF2.PdfWriter()
writer.add_page(reader.pages[1])

output_file = open('archivo_nuevo.pdf','wb') # write binary
writer.write(output_file)
output_file.close()
pdfFile.close()


import docx
d = docx.Document('demo.docx')
print(d.paragraphs[1].text)

print(d.paragraphs[1].runs) # Runs son segmentos de texto dentro de un párrafo. Cada uno empieza cuando cambia el estilo (negritas, por ejemplo)
print(d.paragraphs[0].runs[0].text)

d.paragraphs[1].runs[1].underline = True
d.paragraphs[1].runs[1].text = 'cheverengue'
d.save('demo_cambiado.docx')

d.paragraphs[2].style = 'Normal' # se pueden verificar los estilos en Word con ctr+alt+shift+S
d.save('demo_cambiado.docx')

# Creación de un nuevo documento

docu = docx.Document()
docu.add_paragraph('Me encanta programar con python')
docu.add_paragraph('Este es otro párrafo')
docu.paragraphs[1].add_run('Este es un nuevo run')
docu.paragraphs[1].runs[1].bold = True
docu.save('Nuevo docu.docx')


def texto_documento(archivo):
    doc = docx.Document(archivo)
    texto = []
    for parrafo in doc.paragraphs:
        texto.append(parrafo.text)
    return '\n'.join(texto)

print(texto_documento('Nuevo docu.docx'))